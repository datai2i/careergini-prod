"""
DOCX Generator for CareerGini
Generates editable Word documents for resumes and cover letters.
Mirrors the section order and template style of pdf_generator.py.

Templates supported: professional | executive | fresher
"""
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from typing import Dict, Any, List, Optional
import logging

logger = logging.getLogger(__name__)

# ─────────────────────────────────────────────────────────────────────────────
# Palette
# ─────────────────────────────────────────────────────────────────────────────

PALETTES = {
    "professional": {
        "accent":  RGBColor(0x1A, 0x35, 0x57),   # navy
        "sub":     RGBColor(0x5A, 0x66, 0x77),   # cool gray
        "rule":    RGBColor(0x1A, 0x35, 0x57),
        "body":    RGBColor(0x2D, 0x2D, 0x2D),
        "link":    RGBColor(0x25, 0x63, 0xEB),
    },
    "executive": {
        "accent":  RGBColor(0x1C, 0x1C, 0x1C),   # charcoal
        "sub":     RGBColor(0x4A, 0x55, 0x68),   # slate
        "rule":    RGBColor(0xB8, 0x98, 0x5E),   # platinum-gold
        "body":    RGBColor(0x2D, 0x2D, 0x2D),
        "link":    RGBColor(0x25, 0x63, 0xEB),
    },
    "fresher": {
        "accent":  RGBColor(0x0D, 0x94, 0x88),   # teal
        "sub":     RGBColor(0x47, 0x55, 0x69),   # slate
        "rule":    RGBColor(0x0D, 0x94, 0x88),
        "body":    RGBColor(0x2D, 0x2D, 0x2D),
        "link":    RGBColor(0x25, 0x63, 0xEB),
    },
}

# ─────────────────────────────────────────────────────────────────────────────
# Helpers
# ─────────────────────────────────────────────────────────────────────────────

def _clean(v) -> str:
    return str(v or "").replace("Not specified", "").replace("(Not specified)", "").strip()


def _bullets(exp: dict) -> List[str]:
    raw = exp.get("tailored_bullets") or exp.get("key_achievement") or []
    if isinstance(raw, list):
        return [str(b).strip() for b in raw if str(b).strip()]
    return [b.strip() for b in str(raw).replace(";", "\n").split("\n") if b.strip()]


def _set_para_margins(para, left_pt=0, top_pt=0, bottom_pt=0):
    """Set paragraph spacing in points."""
    pPr = para._p.get_or_add_pPr()
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:before"), str(int(top_pt * 20)))   # 20 twips per pt
    spacing.set(qn("w:after"),  str(int(bottom_pt * 20)))
    pPr.append(spacing)
    if left_pt:
        ind = OxmlElement("w:ind")
        ind.set(qn("w:left"), str(int(left_pt * 20)))
        pPr.append(ind)


def _add_hr(doc: Document, color_rgb: RGBColor, thickness_pt: float = 1.0):
    """Add a horizontal rule paragraph."""
    p = doc.add_paragraph()
    _set_para_margins(p, top_pt=2, bottom_pt=4)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(int(thickness_pt * 8)))  # 1/8 pt units
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), f"{color_rgb[0]:02X}{color_rgb[1]:02X}{color_rgb[2]:02X}")
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p


def _set_run_color(run, rgb: RGBColor):
    run.font.color.rgb = rgb


def _add_name(doc, text, palette, compact):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(22 if compact else 30)
    _set_run_color(run, palette["accent"])
    _set_para_margins(p, bottom_pt=2)
    return p


def _add_title(doc, text, palette, compact):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.italic = True
    run.font.size = Pt(11 if compact else 13)
    _set_run_color(run, palette["sub"])
    _set_para_margins(p, bottom_pt=2)
    return p


def _add_contact(doc, persona, palette, compact):
    parts = []
    for key in ("email", "phone", "location", "linkedin", "portfolio_url"):
        v = _clean(persona.get(key))
        if v:
            parts.append(v)
    if not parts:
        return
    p = doc.add_paragraph()
    run = p.add_run("  ·  ".join(parts))
    run.font.size = Pt(8.5 if compact else 9.5)
    _set_run_color(run, RGBColor(0x66, 0x66, 0x66))
    _set_para_margins(p, bottom_pt=3)


def _add_section_header(doc, label, palette, compact):
    """Bold uppercase section label followed by a colored rule."""
    p = doc.add_paragraph()
    run = p.add_run(label.upper())
    run.bold = True
    run.font.size = Pt(9.5 if compact else 11)
    _set_run_color(run, palette["accent"])
    _set_para_margins(p, top_pt=8 if compact else 14, bottom_pt=2)
    _add_hr(doc, palette["rule"], thickness_pt=1.0 if compact else 1.5)


def _add_body(doc, text, palette, compact, italic=False, bold=False, indent_pt=0):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(9 if compact else 10.5)
    run.italic = italic
    run.bold = bold
    _set_run_color(run, palette["body"])
    _set_para_margins(p, left_pt=indent_pt, bottom_pt=2)
    return p


def _add_bullet(doc, text, palette, compact):
    """
    Creates a manual bullet paragraph. 
    Avoids using 'List Bullet' style which can be missing in default templates.
    """
    p = doc.add_paragraph()
    # Manual indentation for bullet look
    pf = p.paragraph_format
    pf.left_indent = Pt(18 if compact else 24)
    pf.first_line_indent = Pt(-12 if compact else -18)
    
    run = p.add_run(f"•  {text}")
    run.font.size = Pt(9 if compact else 10.5)
    _set_run_color(run, palette["body"])
    _set_para_margins(p, bottom_pt=2)


def _add_experience_block(doc, exp, palette, compact, max_bullets=None):
    role    = _clean(exp.get("role"))
    company = _clean(exp.get("company"))
    dur     = _clean(exp.get("duration"))
    blist   = _bullets(exp)
    if max_bullets:
        blist = blist[:max_bullets]

    if role:
        p = doc.add_paragraph()
        run = p.add_run(role)
        run.bold = True
        run.font.size = Pt(9.5 if compact else 11)
        _set_run_color(run, palette["accent"])
        _set_para_margins(p, top_pt=5 if compact else 8, bottom_pt=1)

    if company or dur:
        p = doc.add_paragraph()
        run = p.add_run("  ·  ".join([x for x in [company, dur] if x]))
        run.italic = True
        run.font.size = Pt(8.5 if compact else 10)
        _set_run_color(run, RGBColor(0x66, 0x66, 0x66))
        _set_para_margins(p, bottom_pt=2)

    for b in blist:
        _add_bullet(doc, b, palette, compact)


def _add_education_block(doc, persona, palette, compact):
    for edu in (persona.get("education") or []):
        d  = _clean(edu.get("degree"))
        sc = _clean(edu.get("school"))
        yr = _clean(edu.get("year"))
        parts = []
        if d:  parts.append(d)
        if sc: parts.append(sc)
        txt = ", ".join(parts)
        if yr: txt += f"  ({yr})"
        if txt.strip():
            p = doc.add_paragraph()
            r_degree = p.add_run(d + (", " if sc else ""))
            r_degree.bold = True
            r_degree.font.size = Pt(9 if compact else 10.5)
            _set_run_color(r_degree, palette["body"])
            if sc:
                r_school = p.add_run(sc + (f"  ({yr})" if yr else ""))
                r_school.font.size = Pt(9 if compact else 10.5)
                _set_run_color(r_school, palette["body"])
            _set_para_margins(p, bottom_pt=3)


def _setup_doc(template: str, page_count: int) -> Document:
    """Create a Document with correct margins."""
    doc = Document()
    compact = page_count == 1
    for section in doc.sections:
        margin = Cm(1.07 if compact else 1.52)   # ≈ 0.42 in / 0.60 in
        section.top_margin    = margin
        section.bottom_margin = margin
        section.left_margin   = margin
        section.right_margin  = margin
    # Remove default paragraph spacing
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10.5)
    pf = style.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after  = Pt(0)
    return doc


# ─────────────────────────────────────────────────────────────────────────────
# Template renderers
# ─────────────────────────────────────────────────────────────────────────────

def _render_professional_docx(doc, persona, palette, compact):
    max_bullets = 2 if compact else None
    skills = (persona.get("top_skills") or [])[:8 if compact else 20]

    _add_name(doc, _clean(persona.get("full_name")), palette, compact)
    _add_title(doc, _clean(persona.get("professional_title")), palette, compact)
    _add_contact(doc, persona, palette, compact)
    _add_hr(doc, palette["accent"], thickness_pt=1.5)

    summary = _clean(persona.get("summary"))
    if summary:
        _add_section_header(doc, "Professional Summary", palette, compact)
        _add_body(doc, summary, palette, compact)

    if skills:
        _add_section_header(doc, "Core Skills", palette, compact)
        _add_body(doc, ", ".join(skills), palette, compact)

    exps = persona.get("experience_highlights") or []
    if exps:
        _add_section_header(doc, "Work Experience", palette, compact)
        for exp in exps:
            _add_experience_block(doc, exp, palette, compact, max_bullets=max_bullets)

    if persona.get("education"):
        _add_section_header(doc, "Education", palette, compact)
        _add_education_block(doc, persona, palette, compact)

    if not compact and persona.get("projects"):
        _add_section_header(doc, "Projects", palette, compact)
        for proj in (persona.get("projects") or []):
            n = _clean(proj.get("name"))
            d = _clean(proj.get("description"))
            if n: _add_body(doc, n, palette, compact, bold=True)
            if d: _add_body(doc, d, palette, compact)

    if not compact and persona.get("certifications"):
        _add_section_header(doc, "Certifications", palette, compact)
        for c in (persona.get("certifications") or []):
            _add_bullet(doc, _clean(c), palette, compact)


def _render_executive_docx(doc, persona, palette, compact):
    max_bullets = 3 if compact else None
    skills = (persona.get("top_skills") or [])[:9 if compact else 12]

    _add_name(doc, _clean(persona.get("full_name")), palette, compact)
    _add_title(doc, _clean(persona.get("professional_title")), palette, compact)
    _add_contact(doc, persona, palette, compact)
    # Double rule for executive
    _add_hr(doc, palette["rule"], thickness_pt=2.5)
    _add_hr(doc, RGBColor(0xCC, 0xCC, 0xCC), thickness_pt=0.5)

    summary = _clean(persona.get("summary"))
    if summary:
        _add_section_header(doc, "Executive Profile", palette, compact)
        _add_body(doc, summary, palette, compact)

    if skills:
        _add_section_header(doc, "Core Competencies", palette, compact)
        # Three per line
        lines = [skills[i:i+3] for i in range(0, len(skills), 3)]
        for line in lines:
            _add_body(doc, "     ·     ".join(line), palette, compact)

    exps = persona.get("experience_highlights") or []
    if exps:
        _add_section_header(doc, "Career Timeline", palette, compact)
        for exp in exps:
            _add_experience_block(doc, exp, palette, compact, max_bullets=max_bullets)

    if persona.get("education"):
        _add_section_header(doc, "Education", palette, compact)
        _add_education_block(doc, persona, palette, compact)

    if not compact and persona.get("projects"):
        _add_section_header(doc, "Key Achievements & Projects", palette, compact)
        for proj in (persona.get("projects") or []):
            n = _clean(proj.get("name"))
            d = _clean(proj.get("description"))
            if n: _add_body(doc, n, palette, compact, bold=True)
            if d: _add_body(doc, d, palette, compact)

    if not compact and persona.get("certifications"):
        _add_section_header(doc, "Certifications & Credentials", palette, compact)
        for c in (persona.get("certifications") or []):
            _add_bullet(doc, _clean(c), palette, compact)


def _render_fresher_docx(doc, persona, palette, compact):
    max_bullets = 2 if compact else 3
    skills = (persona.get("top_skills") or [])[:6 if compact else 12]
    projects = (persona.get("projects") or [])[:2 if compact else None]

    _add_name(doc, _clean(persona.get("full_name")), palette, compact)
    _add_title(doc, _clean(persona.get("professional_title")), palette, compact)
    _add_contact(doc, persona, palette, compact)
    _add_hr(doc, palette["accent"], thickness_pt=1.5)

    summary = _clean(persona.get("summary"))
    if summary:
        _add_section_header(doc, "Career Objective", palette, compact)
        _add_body(doc, summary, palette, compact)

    if persona.get("education"):
        _add_section_header(doc, "Education", palette, compact)
        _add_education_block(doc, persona, palette, compact)

    if skills:
        _add_section_header(doc, "Technical Skills", palette, compact)
        # 3 per line (tag-style in text)
        lines = [skills[i:i + (3 if compact else 4)] for i in range(0, len(skills), 3 if compact else 4)]
        for line in lines:
            _add_body(doc, "  |  ".join(line), palette, compact)

    if projects:
        _add_section_header(doc, "Projects & Academic Work", palette, compact)
        for proj in projects:
            n = _clean(proj.get("name"))
            d = _clean(proj.get("description"))
            if n: _add_body(doc, n, palette, compact, bold=True)
            if d: _add_body(doc, d, palette, compact)

    exps = persona.get("experience_highlights") or []
    if exps:
        _add_section_header(doc, "Work Experience & Internships", palette, compact)
        for exp in exps:
            _add_experience_block(doc, exp, palette, compact, max_bullets=max_bullets)

    if not compact and persona.get("certifications"):
        _add_section_header(doc, "Certifications & Achievements", palette, compact)
        for c in (persona.get("certifications") or []):
            _add_bullet(doc, _clean(c), palette, compact)


# ─────────────────────────────────────────────────────────────────────────────
# Public API
# ─────────────────────────────────────────────────────────────────────────────

def generate_resume_docx(
    output_path: str,
    persona: Dict[str, Any],
    template: str = "professional",
    page_count: int = 2,
) -> bool:
    """Generate an editable DOCX resume."""
    if template not in ("professional", "executive", "fresher"):
        template = "professional"

    compact = (page_count == 1)
    palette = PALETTES[template]

    # Normalise persona
    persona = dict(persona)
    if "tailored_summary"    in persona: persona["summary"]                = persona["tailored_summary"]
    if "tailored_skills"     in persona: persona["top_skills"]             = persona["tailored_skills"]
    if "tailored_experience" in persona: persona["experience_highlights"]  = persona["tailored_experience"]

    try:
        logger.info(f"Starting DOCX generation: {template}, {page_count}p")
        doc = _setup_doc(template, page_count)
        if template == "executive":
            _render_executive_docx(doc, persona, palette, compact)
        elif template == "fresher":
            _render_fresher_docx(doc, persona, palette, compact)
        else:
            _render_professional_docx(doc, persona, palette, compact)
        
        doc.save(output_path)
        logger.info(f"DOCX successfully built and saved → {output_path}")
        return True
    except Exception as e:
        logger.error(f"FATAL: DOCX generation failed for {output_path}. Error: {e}", exc_info=True)
        return False


def generate_cover_letter_docx(
    output_path: str,
    persona: Dict[str, Any],
    template: str = "professional",
) -> bool:
    """Generate an editable DOCX cover letter."""
    if template not in ("professional", "executive", "fresher"):
        template = "professional"

    palette = PALETTES[template]
    compact = False

    try:
        doc = _setup_doc(template, page_count=2)
        logger.info(f"Starting Cover Letter DOCX generation: {template}")

        _add_name(doc, _clean(persona.get("full_name")), palette, compact)
        _add_contact(doc, persona, palette, compact)
        _add_hr(doc, palette["accent"], thickness_pt=1.5)

        cl = _clean(persona.get("cover_letter", ""))
        if cl:
            for para_text in cl.split("\n"):
                if para_text.strip():
                    _add_body(doc, para_text.strip(), palette, compact)
                else:
                    doc.add_paragraph()   # blank line between paragraphs

        doc.save(output_path)
        logger.info(f"Cover Letter DOCX → {output_path}")
        return True
    except Exception as e:
        logger.error(f"Cover Letter DOCX generation failed: {e}")
        return False
